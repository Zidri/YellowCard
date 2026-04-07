from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from PyPDF2 import PdfReader
from docx import Document
import os
import uuid

app = Flask(__name__)
CORS(app)

# --------------------
# Folders
# --------------------
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# --------------------
# Convert PDF to Word (FIXED)
# --------------------
@app.route("/convert", methods=["POST"])
def convert_pdf_to_word():
    try:
        # Check if file exists
        if "file" not in request.files:
            return jsonify({"error": "No file uploaded"}), 400

        file = request.files["file"]

        # Validate filename
        if file.filename == "":
            return jsonify({"error": "Empty filename"}), 400

        if not file.filename.lower().endswith(".pdf"):
            return jsonify({"error": "Invalid file type"}), 400

        # Save PDF with unique name (prevents overwrite)
        unique_pdf_name = f"{uuid.uuid4()}.pdf"
        pdf_path = os.path.join(UPLOAD_FOLDER, unique_pdf_name)
        file.save(pdf_path)

        # Read PDF
        reader = PdfReader(pdf_path)

        if len(reader.pages) == 0:
            return jsonify({"error": "PDF has no pages"}), 400

        # Create Word document
        doc = Document()

        for i, page in enumerate(reader.pages):
            text = page.extract_text()

            if text:
                doc.add_paragraph(text)
            else:
                doc.add_paragraph(f"[No text found on page {i+1}]")

        # Save Word file with unique name
        output_filename = f"{uuid.uuid4()}.docx"
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)
        doc.save(output_path)

        # Return file
        return send_file(output_path, as_attachment=True)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --------------------
# Test route (optional but useful)
# --------------------
@app.route("/", methods=["GET"])
def home():
    return "Backend is running!"

# --------------------
# Run app
# --------------------
if __name__ == "__main__":
    app.run(debug=True, port=5000)